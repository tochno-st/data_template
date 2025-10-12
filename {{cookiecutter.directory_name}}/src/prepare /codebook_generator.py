from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION
import pandas as pd

class CodebookGenerator:
    """Generate codebook for the given dataframe in docx format"""
    def __init__(self, data_path, dataset_name: str, dataset_source: str, dataset_short: str, dataset_num: int):
        self.data_path = data_path
        self.data = pd.read_csv(data_path, sep=';', encoding='utf-8', dtype="string")
        self.codebook = Document()
        self.dataset_name = dataset_name
        self.dataset_source = dataset_source
        self.dataset_short = dataset_short
        self.dataset_num = dataset_num

    @staticmethod
    def add_section_header(doc, title):
        heading = doc.add_paragraph(title, style='Heading1')
        heading.paragraph_format.space_after = Pt(12)

    @staticmethod
    def add_table(doc, headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Add rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

    @staticmethod
    def set_font_for_style(style, font_name):
        font = style.font
        font.name = font_name
        # Ensure Cyrillic support
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)  # ASCII font
        rFonts.set(qn('w:hAnsi'), font_name)  # High ANSI font
        rFonts.set(qn('w:eastAsia'), font_name)  # East Asian font (for Cyrillic)

    def set_heading_font(self, doc, font_name):
        for heading_level in range(1, 7):  # Heading1 to Heading6
            heading_style = doc.styles[f'Heading {heading_level}']
            self.set_font_for_style(heading_style, font_name)

    def generate_codebook(self):
        self.set_font_for_style(self.codebook.styles['Normal'], 'Lora')
        self.set_heading_font(self.codebook, 'Lora')

        ###Section 1###
        p = self.codebook.add_paragraph()
        run = p.add_run(f"«{self.dataset_name}»")
        run.bold = True
        font = run.font
        font.size = Pt(18)
        font.color.rgb = RGBColor(230, 86, 57)

        self.add_section_header(self.codebook, "1. История изменений")

        table = self.codebook.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        column_widths = [Cm(3.0), Cm(3.0), Cm(10.0)]

        for i, width in enumerate(column_widths):
            col = table.columns[i]
            for cell in col.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width.twips)))  # Convert centimeters to twentieths of a point (twips)
                tcW.set(qn('w:type'), 'dxa')  # Use 'dxa' for twentieths of a point
                tcPr.append(tcW)

        # Set table headers
        headers = ["Дата", "Версия", "Описание изменений"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            # Make headers bold
            cell.paragraphs[0].runs[0].bold = True
            # Set background color for the header row
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'f8b7a9')  # Light gray color (hex code)
            cell._element.get_or_add_tcPr().append(shading_elm)

        data = [
            [datetime.now().strftime("%d.%m.%Y"), "1.0", "Документ создан"],
        ]

        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, cell_data in enumerate(row_data):
                table.cell(row_idx, col_idx).text = cell_data

        ###Section 2###

        self.add_section_header(self.codebook, "2. Основные сведения")

        table = self.codebook.add_table(rows=19, cols=2)
        table.style = 'Table Grid'

        column_widths = [Cm(5.0), Cm(12.0)]

        for i, width in enumerate(column_widths):
            col = table.columns[i]
            for cell in col.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width.twips)))  # Convert centimeters to twentieths of a point (twips)
                tcW.set(qn('w:type'), 'dxa')  # Use 'dxa' for twentieths of a point
                tcPr.append(tcW)

        headers = ["Атрибут", "Значение"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'f8b7a9')  # Light gray color (hex code)
            cell._element.get_or_add_tcPr().append(shading_elm)


        second_row = table.rows[1]
        second_row.cells[0].merge(second_row.cells[1])  # Merge the two cells in the second row
        cell = second_row.cells[0]
        cell.text = "Общие сведения"  # Add the merged cell value
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.runs[0].bold = True

        data = [["Наименование", self.dataset_name], 
                ["Краткое описание","" ],
                [ "Тематика",""],
                ["Единица наблюдения",""],
                ["Количество атрибутов",str(len(self.data.columns))],
                ["Количество наблюдений",str(len(self.data)-1)],
                ["Доступные форматы",""]]


        for row_idx, row_data in enumerate(data, start=2):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data

        ten_row = table.rows[9]
        ten_row.cells[0].merge(ten_row.cells[1]) 
        cell = ten_row.cells[0]
        cell.text = "Периодичность публикации и обновления в каталоге «Если быть точным»"  # Add the merged cell value
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.runs[0].bold = True

        data = [["Дата размещения в каталоге",""],
                ["Покрываемый временной период",""],
                ["Дата последнего обновления набора данных",""]]

        for row_idx, row_data in enumerate(data, start=2):  # Start from the third row
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx+8, col_idx)
                cell.text = cell_data


        ten_row = table.rows[13]
        ten_row.cells[0].merge(ten_row.cells[1]) 
        cell = ten_row.cells[0]
        cell.text = "Дополнительные сведения"  # Add the merged cell value
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.runs[0].bold = True


        data = [["Цитирование набора данных на русском языке",
                f"«{self.dataset_name}» // {self.dataset_source}; обработка: «Если быть точным», 2025. URL:  https://tochno.st/{self.dataset_short}"],
                ["For references (English)",
                 f"«{self.dataset_name}» // {self.dataset_source}; data-processing: «To Be Precise», 2025. URL:  https://tochno.st/{self.dataset_short}"],
                ["Ссылка на открытый репозиторий", "Открытого репозитория нет"],
                ["Геоданные",""],
                ["Лицензия", "Creative Commons BY"]]

        for row_idx, row_data in enumerate(data, start=2):  # Start from the third row
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx+12, col_idx)
                cell.text = cell_data

        self.codebook.add_page_break()

        ###Section 3###

        section = self.codebook.add_section()
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Inches(11)  # Landscape width
        section.page_height = Inches(8.5)  # Landscape height

        self.add_section_header(self.codebook, "3. Структура набора данных")

        columns_number = self.data.shape[1]

        table = self.codebook.add_table(rows=columns_number+1, cols=6)
        table.style = 'Table Grid'

        column_widths = [Cm(1.0), Cm(3.0), Cm(9.0), Cm(3.0), Cm(3.0), Cm(3.0)]

        for i, width in enumerate(column_widths):
            col = table.columns[i]
            for cell in col.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width.twips)))  # Convert centimeters to twentieths of a point (twips)
                tcW.set(qn('w:type'), 'dxa')  # Use 'dxa' for twentieths of a point
                tcPr.append(tcW)

        headers = ["№", "Атрибут", "Описание", "Число пропусков", "Единица измерения", "Тип данных"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            # Make headers bold
            cell.paragraphs[0].runs[0].bold = True
            # Set background color for the header row
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'f8b7a9')  # Light gray color (hex code)
            cell._element.get_or_add_tcPr().append(shading_elm)


        # Create data for codebook table
        data = []
        for i, col in enumerate(self.data.columns, 1):
            missing_count = self.data[col].isna().sum() + (self.data[col] == "ND").sum()
            data.append([str(i), col, "", str(missing_count), "", ""])

        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx + 1, col_idx)  # +1 to skip header row
                cell.text = cell_data
                if col_idx == 1:  # Second column (attribute names)
                    cell.paragraphs[0].runs[0].bold = True
                if col_idx == 3:  # Fourth column (missing count)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        self.codebook.add_page_break()

        ###Section 4###

        section = self.codebook.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT  # Reset to portrait
        section.page_width = Inches(8.5)  # Portrait width
        section.page_height = Inches(11)  # Portrait height

        self.add_section_header(self.codebook, "4. Известные ограничения")
        p = self.codebook.add_paragraph("Полнота данных")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.bold = True
        run.italic = True

        p = self.codebook.add_paragraph("Сопоставимость данных")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.bold = True
        run.italic = True

        self.codebook.add_page_break()

        ###Section 5###

        self.add_section_header(self.codebook, "5. Источники")

        table = self.codebook.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        column_widths = [Cm(5.0), Cm(12.0)]

        for i, width in enumerate(column_widths):
            col = table.columns[i]
            for cell in col.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width.twips)))  # Convert centimeters to twentieths of a point (twips)
                tcW.set(qn('w:type'), 'dxa')  # Use 'dxa' for twentieths of a point
                tcPr.append(tcW)

        headers = ["Атрибут", "Значение"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            # Make headers bold
            cell.paragraphs[0].runs[0].bold = True
            # Set background color for the header row
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'f8b7a9')  # Light gray color (hex code)
            cell._element.get_or_add_tcPr().append(shading_elm)

        data = [["Наименование источника данных", ""],
                ["Владелец(ы) источника данных", ""],
                ["Краткое описание источника данных", ""],
                ["Ссылка на источник данных", "" ]]

        for row_idx, row_data in enumerate(data, start=1):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_data
                
        for para in self.codebook.paragraphs:
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name[-1])
                
                if heading_level == 1:
                    for run in para.runs:
                        font = run.font
                        font.name = 'Lora'
                        font.size = Pt(14)
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                elif heading_level == 2:
                    for run in para.runs:
                        font = run.font
                        font.name = 'Lora'
                        font.size = Pt(14)
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        cur_date = datetime.now().strftime('%Y%m%d')

        return self.codebook, f"description_{self.dataset_short}_{self.dataset_num}_v{cur_date}.docx"
